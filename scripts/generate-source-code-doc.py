#!/usr/bin/env python3
"""Generate RIDE Source Code Documentation as a Word (.docx) file."""

from __future__ import annotations

import os
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "RIDE_Source_Code_Documentation.docx"

SOURCE_EXTENSIONS = {
    ".php", ".sql", ".js", ".css", ".xml", ".htaccess", ".md", ".gitignore"
}

EXCLUDE_DIRS = {
    ".git", "node_modules", "vendor", "__pycache__",
}

EXCLUDE_PATH_PARTS = {
    "public/assets/uploads",
    "storage/avatars",
    "storage/scholarly",
    "storage/signatures",
}

MODULES: list[tuple[str, str, list[str]]] = [
    (
        "1. Core Application",
        "Bootstrap, configuration, helpers, and routing entry points.",
        [
            "app/bootstrap.php",
            "app/helpers.php",
            "app/config/",
            "public/index.php",
            "routes/web.php",
            "routes/api.php",
        ],
    ),
    (
        "2. Core Framework",
        "Authentication, database, and HTTP router.",
        [
            "app/Core/",
        ],
    ),
    (
        "3. Middleware",
        "Request authentication and role-based access control.",
        [
            "app/Middleware/",
        ],
    ),
    (
        "4. Controllers",
        "HTTP request handlers for all application modules.",
        [
            "app/Controllers/",
        ],
    ),
    (
        "5. Models",
        "Database entity models and data access layer.",
        [
            "app/Models/",
        ],
    ),
    (
        "6. Services",
        "Business logic services for consolidation and access control.",
        [
            "app/Services/",
        ],
    ),
    (
        "7. Support Classes",
        "Shared utility and helper classes.",
        [
            "app/Support/",
        ],
    ),
    (
        "8. Authentication Module",
        "Login, registration, and session management views.",
        [
            "app/views/auth/",
        ],
    ),
    (
        "9. Dashboard Module",
        "Main dashboard, analytics, and messaging hub.",
        [
            "app/views/dashboard/",
        ],
    ),
    (
        "10. Proposals Module",
        "Research proposal forms, workflow, and consolidated reports.",
        [
            "app/views/proposals/",
        ],
    ),
    (
        "11. Projects Module",
        "Approved project monitoring, milestones, and reports.",
        [
            "app/views/projects/",
        ],
    ),
    (
        "12. Monitoring Module",
        "System-wide monitoring and oversight views.",
        [
            "app/views/monitoring/",
        ],
    ),
    (
        "13. Scholarly Module",
        "Scholarly output tracking and attachments.",
        [
            "app/views/scholarly/",
        ],
    ),
    (
        "14. Admin Module",
        "Account management, highlights, and scholarly admin.",
        [
            "app/views/admin/",
        ],
    ),
    (
        "15. Profile Module",
        "User profile, avatar, and signature management.",
        [
            "app/views/profile/",
        ],
    ),
    (
        "16. Messages Module",
        "Direct messaging between users.",
        [
            "app/views/messages/",
        ],
    ),
    (
        "17. Co-Author Invitations Module",
        "Co-author invitation workflow views.",
        [
            "app/views/coauthor-invitations/",
        ],
    ),
    (
        "18. Reports Module",
        "Accreditation and extension beneficiary reports.",
        [
            "app/views/reports/",
        ],
    ),
    (
        "19. Settings Module",
        "Faculty deadline and system settings.",
        [
            "app/views/settings/",
        ],
    ),
    (
        "20. Layouts & Partials",
        "Shared layout templates and reusable partials.",
        [
            "app/views/layouts/",
            "app/views/partials/",
            "app/views/errors/",
        ],
    ),
    (
        "21. Frontend Assets",
        "CSS, JavaScript, and public web configuration.",
        [
            "public/assets/css/",
            "public/assets/js/",
            "public/.htaccess",
        ],
    ),
    (
        "22. Database",
        "Schema, seeds, and migration scripts.",
        [
            "database/",
        ],
    ),
    (
        "23. Scripts",
        "Installation, seeding, and maintenance utilities.",
        [
            "scripts/",
        ],
    ),
    (
        "24. Java REST API Client",
        "Maven-based Java client for external API integration.",
        [
            "java-client/",
        ],
    ),
    (
        "25. Documentation",
        "Project README and documentation.",
        [
            "README.md",
        ],
    ),
]


def normalize_path(path: Path) -> str:
    return str(path.relative_to(ROOT)).replace("\\", "/")


def is_source_file(path: Path) -> bool:
    rel = normalize_path(path)
    for part in EXCLUDE_PATH_PARTS:
        if part in rel:
            return False
    if path.name == ".gitkeep":
        return False
    if path.suffix.lower() in SOURCE_EXTENSIONS:
        return True
    if path.name == ".htaccess":
        return True
    return False


def collect_all_source_files() -> list[Path]:
    files: list[Path] = []
    for dirpath, dirnames, filenames in os.walk(ROOT):
        dirnames[:] = [d for d in dirnames if d not in EXCLUDE_DIRS]
        for name in filenames:
            path = Path(dirpath) / name
            if is_source_file(path):
                files.append(path)
    return sorted(files, key=lambda p: normalize_path(p).lower())


def match_module(rel_path: str, patterns: list[str]) -> bool:
    for pattern in patterns:
        pattern = pattern.replace("\\", "/").rstrip("/")
        if pattern.endswith("/"):
            if rel_path.startswith(pattern):
                return True
        elif rel_path == pattern or rel_path.startswith(pattern + "/"):
            return True
    return False


def assign_files_to_modules(all_files: list[Path]) -> dict[str, list[Path]]:
    assigned: dict[str, list[Path]] = {title: [] for title, _, _ in MODULES}
    unassigned: list[Path] = []

    for path in all_files:
        rel = normalize_path(path)
        matched = False
        for title, _, patterns in MODULES:
            if match_module(rel, patterns):
                assigned[title].append(path)
                matched = True
                break
        if not matched:
            unassigned.append(path)

    if unassigned:
        assigned["26. Other Files"] = unassigned

    return assigned


def add_title_page(doc: Document) -> None:
    title = doc.add_heading("RIDE Integrated Monitoring & Workflow Approval System", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Complete Source Code Documentation")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"\nGenerated: {datetime.now().strftime('%B %d, %Y at %H:%M')}\n").font.size = Pt(11)
    meta.add_run(f"Project Path: {ROOT}\n").font.size = Pt(10)
    meta.add_run("\nWestern Philippines University — Research, Innovation, Development & Extension (RIDE)\n").font.size = Pt(11)

    doc.add_page_break()


def add_table_of_contents(doc: Document, module_files: dict[str, list[Path]]) -> None:
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph(
        "This document contains the complete source code for the RIDE system, "
        "organized by functional module. Each section lists all source files "
        "followed by their full source code."
    )

    total_files = 0
    total_lines = 0

    for title in module_files:
        files = module_files[title]
        if not files:
            continue
        lines = sum(count_lines(f) for f in files)
        total_files += len(files)
        total_lines += lines
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(f"{title} ({len(files)} files, {lines:,} lines)")
        run.bold = True

    doc.add_paragraph()
    summary = doc.add_paragraph()
    summary.add_run(f"Total: {total_files} source files, {total_lines:,} lines of code").bold = True
    doc.add_page_break()


def count_lines(path: Path) -> int:
    try:
        return len(path.read_text(encoding="utf-8", errors="replace").splitlines())
    except OSError:
        return 0


def read_source(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def add_module_section(doc: Document, title: str, description: str, files: list[Path]) -> None:
    if not files:
        return

    doc.add_heading(title, level=1)
    doc.add_paragraph(description)

    file_list = doc.add_paragraph()
    file_list.add_run("Source Files:").bold = True
    for path in files:
        rel = normalize_path(path)
        lines = count_lines(path)
        item = doc.add_paragraph(style="List Bullet")
        item.add_run(f"{rel} ({lines} lines)")

    doc.add_paragraph()

    for path in files:
        rel = normalize_path(path)
        doc.add_heading(rel, level=2)

        content = read_source(path)
        code_para = doc.add_paragraph()
        code_run = code_para.add_run(content)
        code_run.font.name = "Consolas"
        code_run.font.size = Pt(8)

        doc.add_paragraph()


def generate_document() -> Path:
    all_files = collect_all_source_files()
    module_files = assign_files_to_modules(all_files)

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title_page(doc)
    add_table_of_contents(doc, module_files)

    module_descriptions = {title: desc for title, desc, _ in MODULES}

    for title, _, _ in MODULES:
        files = module_files.get(title, [])
        if files:
            add_module_section(doc, title, module_descriptions[title], files)

    if "26. Other Files" in module_files:
        add_module_section(
            doc,
            "26. Other Files",
            "Additional source files not assigned to a specific module.",
            module_files["26. Other Files"],
        )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = generate_document()
    all_files = collect_all_source_files()
    print(f"Generated: {output}")
    print(f"Total source files: {len(all_files)}")
    print(f"File size: {output.stat().st_size / 1024 / 1024:.2f} MB")
